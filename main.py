import docx
import random


mydoc = docx.Document()
folder_path = "C:\\Users\\oxana\\Desktop\\Liza.docx"  # path to folder and name of file


def save_message(name):  # method for save to file and send message to console
    mydoc.save (folder_path)
    print(f'{name}')


def calc_chr():  # method for randomly generate which math action will do
    calc_num = random.randint(0, 1)
    if calc_num == 0:
        calcul_char = '-'
    else:
        calcul_char = '+'
    return calcul_char


def rnd_num():  # method for randomly generate numbers for example
    number = random.randint(1, 9)  # <---- change range of numbers
    return number


def calc_exemp ():  # method for checking numbers and totals for example
    number1 = rnd_num()
    number2 = rnd_num()
    calc_char = calc_chr()
    if calc_char == '-':
        while number1 - number2 < 0: # check if total below zero
            number1 = rnd_num()
            number2 = rnd_num()
    else:
        while number1 + number2 > 10: # check if total above ten
            number1 = rnd_num()
            number2 = rnd_num()
    exemp_str = str(str(number1) + str(calc_char) + str(number2) + ' = [\t\t]')  # generate string of one example
    return exemp_str


if __name__ == '__main__':
    mydoc.add_paragraph('TIME 1: [\t\t\t]\t\tTIME 2: [\t\t\t]')
    save_message('File is .... created!')
    for i in range (1, 25):  # iterations for putting strings to .docx file. (1,25) - 25 strings with examples
        calc_str1 = calc_exemp()
        calc_str2 = calc_exemp()
        calc_str3 = calc_exemp()
        calc_str4 = calc_exemp()
        calc_str = str(str(calc_str1) + '\t' + str(calc_str2) + '\t' + str(calc_str3) + '\t' + str(calc_str4))
        mydoc.add_paragraph(calc_str)
    save_message('File is .... filled!')

